import re
import json
import requests
import base64
import pandas as pd
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.conf import settings
from typing import Dict, Any, List

# API URL from settings
API_URL = settings.API_URL

def extract_analysis_context(content: str) -> str:
    """Extract just the context section from analysis content."""
    # Look for "Contexto de Aplicação" section
    context_match = re.search(r'Contexto\s+de\s+Aplicação:?\s*([^\n]+(?:\n(?!\n)[^\n]+)*)', content)
    if context_match:
        return context_match.group(1).strip()
    
    # Alternative pattern if the above doesn't work
    context_match = re.search(r'Contexto:?\s*([^\n]+(?:\n(?!\n)[^\n]+)*)', content)
    if context_match:
        return context_match.group(1).strip()
    
    # If no context section is found, return the content without the analysis title and variables
    # First remove the title line
    content = re.sub(r'^Análise\s+\d+\s*[-:]\s*[^\n]+', '', content, flags=re.MULTILINE)
    # Then remove variables section
    content = re.sub(r'\*\*Variáveis\s+a\s+[Ss]erem\s+[Uu]tilizadas\*\*[\s\S]*?Variáveis\s+dependentes:.*[\s\S]*?Variáveis\s+independentes:.*?(?=\n\n|\Z)', '', content)
    # Clean up and return
    return content.strip()

def upload_file(file) -> Dict[str, Any]:
    """Upload a file to the API."""
    try:
        files = {"file": (file.name, file, "multipart/form-data")}
        response = requests.post(f"{API_URL}/api/data/upload", files=files)
        return response.json()
    except Exception as e:
        return {"success": False, "message": str(e)}

def get_dataset_preview(upload_id: str, rows: int = 10) -> Dict[str, Any]:
    """Get dataset preview."""
    try:
        response = requests.get(f"{API_URL}/api/data/preview/{upload_id}?rows={rows}")
        return response.json()
    except Exception as e:
        return {"success": False, "message": str(e)}

def get_treated_dataset_preview(upload_id: str, rows: int = 10) -> Dict[str, Any]:
    """Get updated dataset preview after treatment."""
    try:
        response = requests.get(f"{API_URL}/api/data-treatment/preview/{upload_id}?rows={rows}")
        if response.status_code == 200:
            return response.json()
        else:
            return {"success": False, "message": "Erro ao obter prévia dos dados"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def request_recommendations(upload_id: str, dependent_vars: List[str], 
                           independent_vars: List[str], 
                           column_descriptions: Dict[str, str],
                           objective: str) -> Dict[str, Any]:
    """Request analysis recommendations from the API."""
    try:
        payload = {
            "upload_id": upload_id,
            "variables": {
                "dependent_vars": dependent_vars,
                "independent_vars": independent_vars,
                "column_descriptions": column_descriptions,
                "objective": objective
            }
        }
        response = requests.post(f"{API_URL}/api/analysis/recommend", json=payload)
        return response.json()
    except Exception as e:
        return {"success": False, "message": str(e)}

def execute_analysis(upload_id: str, analysis_number: int, analysis_detail: str, dep_vars: List, indep_vars: List) -> Dict[str, Any]:
    """Execute a specific analysis."""
    try:                
        # Prepare payload with extracted variables if available        
        payload = {
            "upload_id": upload_id,
            "analysis_number": int(analysis_number),
            "analysis_detail": analysis_detail
        }
        
        # Always add extracted variables to the payload to preserve them
        if dep_vars:
            payload["dependent_vars"] = dep_vars
        if indep_vars:
            payload["independent_vars"] = indep_vars
        
        # Make the execution request directly
        response = requests.post(f"{API_URL}/api/analysis/execute", json=payload, timeout=120)
        
        # Check if the response was successful
        if response.status_code != 200:
            return {
                "success": False,
                "message": f"Erro API ({response.status_code}): {response.text}"
            }
            
        return response.json()
        
    except Exception as e:
        return {"success": False, "message": str(e)}

def execute_all_analyses(upload_id: str) -> Dict[str, Any]:
    """Execute all analyses."""
    try:
        response = requests.post(f"{API_URL}/api/analysis/execute-all/{upload_id}")
        return response.json()
    except Exception as e:
        return {"success": False, "message": str(e)}

def get_analyses_status(upload_id: str) -> Dict[str, Any]:
    """Get the status of analyses for a dataset."""
    try:
        response = requests.get(f"{API_URL}/api/analysis/status/{upload_id}")
        return response.json()
    except Exception as e:
        return {"success": False, "message": str(e)}

def generate_html_report(analyses):
    """Generates an HTML consolidated report of all analyses."""
    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Relatório de Análises - ChatSmart</title>
        <style>
            body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 1200px; margin: 0 auto; padding: 20px; }
            h1 { color: #0066cc; border-bottom: 2px solid #99ccff; padding-bottom: 10px; }
            h2 { color: #0066cc; margin-top: 30px; }
            h3 { color: #333; background-color: #f5f5f5; padding: 10px; }
            .analysis-card { border: 1px solid #ddd; border-radius: 5px; margin-bottom: 30px; padding: 20px; }
            .analysis-header { background-color: #f0f7ff; padding: 10px; margin-bottom: 15px; border-radius: 5px; }
            .success { border-left: 5px solid #00cc66; }
            .error { border-left: 5px solid #ff3333; }
            .interpretation { background-color: #f9f9f9; padding: 15px; border-left: 3px solid #0066cc; margin: 10px 0; }
            .visualization { margin: 20px 0; text-align: center; }
            .table-container { overflow-x: auto; margin: 20px 0; }
            footer { margin-top: 50px; text-align: center; color: #666; border-top: 1px solid #eee; padding-top: 10px; }
        </style>
    </head>
    <body>
        <h1>Relatório de Análises - ChatSmart</h1>
        <p>Relatório gerado em """ + pd.Timestamp.now().strftime("%d/%m/%Y %H:%M:%S") + """</p>
    """
    
    # Métricas gerais
    total_analyses = len(analyses)
    completed_analyses = sum(1 for a in analyses if isinstance(a, dict) and a.get('results') and a.get('results', {}).get('success', False))
    failed_analyses = sum(1 for a in analyses if isinstance(a, dict) and a.get('results') and not a.get('results', {}).get('success', False))
    
    html += f"""
        <h2>Visão Geral</h2>
        <ul>
            <li><strong>Total de Análises:</strong> {total_analyses}</li>
            <li><strong>Análises Concluídas:</strong> {completed_analyses}</li>
            <li><strong>Análises com Erro:</strong> {failed_analyses}</li>
        </ul>
    """
    
    # Adicionar cada análise ao relatório
    html += "<h2>Análises</h2>"
    
    for analysis in analyses:
        if not analysis.get('results'):
            continue
            
        result = analysis.get('results', {})
        success = result.get('success', False)
        
        html += f"""
        <div class="analysis-card {'success' if success else 'error'}">
            <div class="analysis-header">
                <h3>Análise {analysis.get('number')}: {analysis.get('name')}</h3>
                <p><strong>Status:</strong> {'✅ Sucesso' if success else '❌ Erro'}</p>
            </div>
            
            <p><strong>Variáveis Dependentes:</strong> {', '.join(analysis.get('dependent_vars', []))}</p>
            <p><strong>Variáveis Independentes:</strong> {', '.join(analysis.get('independent_vars', []))}</p>
        """
        
        if success:
            # Adicionar interpretação
            if result.get('interpretation'):
                html_result = result.get('interpretation').replace('\n', '<br>')
                html += f"""
                <div class="interpretation">
                    <h4>Interpretação</h4>
                    <p>{html_result}</p>
                </div>
                """
                
            # Adicionar visualizações
            if result.get('plots'):
                html += "<h4>Visualizações</h4>"
                for plot in result.get('plots', []):
                    html += f"""
                    <div class="visualization">
                        <img src="data:image/png;base64,{plot.get('data')}" alt="Visualização" style="max-width: 100%; height: auto;">
                    </div>
                    """
                    
            # Adicionar tabelas
            if result.get('tables'):
                html += "<h4>Tabelas</h4>"
                for table in result.get('tables', []):
                    html += f"""
                    <div class="table-container">
                        <p><strong>{table.get('description', 'Tabela')}</strong></p>
                        {table.get('html', '')}
                    </div>
                    """
                    
            # Adicionar saída de texto
            if result.get('text_output'):
                html += f"""
                <h4>Saída de Texto</h4>
                <pre>{result.get('text_output', '')}</pre>
                """
        else:
            # Mostrar mensagem de erro
            html += f"""
            <div class="error-message">
                <h4>Erro</h4>
                <pre>{result.get('error', 'Erro desconhecido')}</pre>
            </div>
            """
            
        html += "</div>"
    
    # Footer
    html += """
        <footer>
            <p>Gerado por ChatSmart - Análises Lean Six Sigma com IA</p>
        </footer>
    </body>
    </html>
    """
    
    return html

def generate_docx_report(analyses):
    """Generates a DOCX consolidated report of all analyses."""
    # Criar um novo documento
    doc = Document()
    
    # Adicionar título e cabeçalho
    doc.add_heading('Relatório de Análises - ChatSmart', 0)
    doc.add_paragraph(f'Relatório gerado em {pd.Timestamp.now().strftime("%d/%m/%Y %H:%M:%S")}')
    
    # Métricas gerais
    doc.add_heading('Visão Geral', level=1)
    total_analyses = len(analyses)
    completed_analyses = sum(1 for a in analyses if isinstance(a, dict) and a.get('results') and a.get('results', {}).get('success', False))
    failed_analyses = sum(1 for a in analyses if isinstance(a, dict) and a.get('results') and not a.get('results', {}).get('success', False))
    
    metrics_para = doc.add_paragraph()
    metrics_para.add_run('Total de Análises: ').bold = True
    metrics_para.add_run(f"{total_analyses}\n")
    metrics_para.add_run('Análises Concluídas: ').bold = True
    metrics_para.add_run(f"{completed_analyses}\n")
    metrics_para.add_run('Análises com Erro: ').bold = True
    metrics_para.add_run(f"{failed_analyses}\n")
    
    # Adicionar cada análise ao relatório
    doc.add_heading('Análises', level=1)
    
    for analysis in analyses:
        if not isinstance(analysis, dict) or not analysis.get('results'):
            continue
            
        result = analysis.get('results', {})
        success = result.get('success', False)
        
        # Adicionar cabeçalho da análise
        doc.add_heading(f"Análise {analysis.get('number')}: {analysis.get('name')}", level=2)
        
        status_para = doc.add_paragraph()
        status_para.add_run('Status: ').bold = True
        status_para.add_run('✅ Sucesso' if success else '❌ Erro')
        
        # Adicionar detalhes das variáveis
        vars_para = doc.add_paragraph()
        vars_para.add_run('Variáveis Dependentes: ').bold = True
        vars_para.add_run(f"{', '.join(analysis.get('dependent_vars', []))}\n")
        vars_para.add_run('Variáveis Independentes: ').bold = True
        vars_para.add_run(f"{', '.join(analysis.get('independent_vars', []))}")
        
        if success:
            # Adicionar interpretação
            if result.get('interpretation'):
                doc.add_heading('Interpretação', level=3)
                doc.add_paragraph(result.get('interpretation'))
                
            # Mencionar visualizações
            if result.get('plots'):
                doc.add_heading('Visualizações', level=3)
                doc.add_paragraph(f"Esta análise contém {len(result.get('plots', []))} visualizações.")
                
                # Não podemos adicionar imagens diretamente do base64, então apenas mencionamos
                doc.add_paragraph("As visualizações estão disponíveis na versão HTML do relatório.")
                
            # Mencionar tabelas
            if result.get('tables'):
                doc.add_heading('Tabelas', level=3)
                for table in result.get('tables', []):
                    doc.add_paragraph(f"- {table.get('description', 'Tabela')}")
                
            # Adicionar saída de texto
            if result.get('text_output'):
                doc.add_heading('Saída de Texto', level=3)
                doc.add_paragraph(result.get('text_output'))
        else:
            # Mostrar mensagem de erro
            doc.add_heading('Erro', level=3)
            doc.add_paragraph(result.get('error', 'Erro desconhecido'))
        
        # Adicionar separador entre análises
        doc.add_paragraph('─' * 50)
    
    # Adicionar rodapé
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = "Gerado por ChatSmart - Análises Lean Six Sigma com IA"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar o documento em um buffer de memória
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def get_ai_comparison_insights(analysis1, analysis2):
    """Get AI-generated insights comparing two analyses."""
    try:
        # Format the analyses data for the API
        analysis1_data = {
            "id": analysis1.id,
            "name": analysis1.name,
            "dependent_vars": analysis1.dependent_vars,
            "independent_vars": analysis1.independent_vars,
            "content": analysis1.content,
            "status": analysis1.status,
            "results": analysis1.results or {}
        }
        
        analysis2_data = {
            "id": analysis2.id,
            "name": analysis2.name,
            "dependent_vars": analysis2.dependent_vars,
            "independent_vars": analysis2.independent_vars,
            "content": analysis2.content,
            "status": analysis2.status,
            "results": analysis2.results or {}
        }
        
        # Call the API endpoint for AI insights
        payload = {
            "analysis1": analysis1_data,
            "analysis2": analysis2_data,
        }
        
        response = requests.post(f"{API_URL}/api/comparison/ai-insights", json=payload)
        
        if response.status_code == 200:
            return response.json()
        else:
            return {
                "insights": "Não foi possível gerar insights de comparação.",
                "improvement_suggestions": "Erro na chamada de API."
            }
    except Exception as e:
        return {
            "insights": f"Erro ao gerar insights: {str(e)}",
            "improvement_suggestions": "Tente novamente mais tarde."
        }
